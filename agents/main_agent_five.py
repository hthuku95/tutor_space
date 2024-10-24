import os
import time
import requests
import logging
from typing import Tuple, Dict
from langchain_community.vectorstores import FAISS
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from docx import Document
from dotenv import load_dotenv
import tiktoken
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_core.runnables import RunnableSequence
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import StrOutputParser

# Load environment variables
load_dotenv()

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Check for OpenAI API key
if not os.getenv("OPENAI_API_KEY"):
    raise ValueError("Missing OPENAI_API_KEY environment variable")

if not os.getenv("LANGCHAIN_API_KEY"):
    raise ValueError("Missing LANGCHAIN_API_KEY environment variable")

# Check for StealthGPT API key
STEALTHGPT_API_KEY = os.getenv("STEALTHGPT_API_KEY")
if not STEALTHGPT_API_KEY:
    raise ValueError("Missing STEALTHGPT_API_KEY environment variable")

# Initialize the model
model = ChatOpenAI(model_name="gpt-4")

def count_tokens(text: str) -> int:
    encoding = tiktoken.encoding_for_model("gpt-4")
    tokens = encoding.encode(text)
    return len(tokens)

def generate_academic_content(question: str, word_count: int, enhanced_prompt: str, chunk_size: int = 750) -> str:
    prompt_template = ChatPromptTemplate.from_messages([
        ("system", enhanced_prompt),
        MessagesPlaceholder(variable_name="chat_history"),
        ("human", "{question}\n\nPrevious content:\n{content_snippet}\n\n{additional_instructions}"),
    ])

    chain = RunnableSequence(
        prompt_template,
        model,
        StrOutputParser(),
    )

    content = ""
    total_words_generated = 0
    words_per_chunk = chunk_size
    num_chunks = word_count // words_per_chunk + (1 if word_count % words_per_chunk else 0)

    logging.info(f"Generating content in {num_chunks} chunks.")

    chat_history = []

    # Generate the introduction
    additional_instructions = "Begin with an introduction for the paper."
    response = chain.invoke({
        "question": question,
        "additional_instructions": additional_instructions,
        "chat_history": chat_history,
        "content_snippet": ""  # Empty for the introduction
    })

    introduction = response.strip()
    content += introduction + "\n\n"
    total_words_generated += len(introduction.split())
    chat_history.append(("human", question))
    chat_history.append(("ai", introduction))
    logging.info(f"Introduction generated. Words so far: {total_words_generated}/{word_count}")

    # Generate the body in chunks
    for chunk_num in range(1, num_chunks + 1):
        remaining_words = word_count - total_words_generated
        if remaining_words <= 0:
            break

        content_snippet = '\n'.join(content.split('\n')[-50:])

        additional_instructions = """
        Continue writing the body of the paper. Pick up from where the previous section ended and maintain the flow of ideas. 
        Cover all necessary points to address the topic comprehensively.
        """

        response = chain.invoke({
            "question": question,
            "additional_instructions": additional_instructions,
            "chat_history": chat_history,
            "content_snippet": content_snippet
        })

        chunk_text = response.strip()
        content += chunk_text + "\n\n"
        words_in_chunk = len(chunk_text.split())
        total_words_generated += words_in_chunk
        chat_history.append(("human", additional_instructions))
        chat_history.append(("ai", chunk_text))
        logging.info(f"Chunk {chunk_num} generated. Words so far: {total_words_generated}/{word_count}")

    # Generate the conclusion
    content_snippet = '\n'.join(content.split('\n')[-50:])

    additional_instructions = """
    Finally, write a conclusion for the paper, summarizing the main points and providing final insights.
    """

    response = chain.invoke({
        "question": question,
        "additional_instructions": additional_instructions,
        "chat_history": chat_history,
        "content_snippet": content_snippet
    })

    conclusion = response.strip()
    content += conclusion
    total_words_generated += len(conclusion.split())
    logging.info(f"Conclusion generated. Total words: {total_words_generated}/{word_count}")

    return content

def store_content_in_vector_store(content: str):
    embeddings = OpenAIEmbeddings()
    doc_chunks = content.split("\n\n")
    vector_store = FAISS.from_texts(doc_chunks, embeddings)
    return vector_store

def humanize_content(text: str, readability: str, purpose: str) -> str:
    max_chars_per_request = 5000
    text_chunks = [text[i:i+max_chars_per_request] for i in range(0, len(text), max_chars_per_request)]
    rephrased_text = ""

    for idx, chunk in enumerate(text_chunks):
        url = "https://stealthgpt.ai/api/stealthify" 
        headers = {
            "api-token": STEALTHGPT_API_KEY,
            "Content-Type": "application/json"
        }
        payload = {
            "prompt": chunk,
            "rephrase": True,
            "tone": readability,
            "mode": "High" 
        }
        logging.info(f"Stealthifying chunk {idx+1}/{len(text_chunks)}. Chunk size: {len(chunk)} characters")
        try:
            response = requests.post(url, headers=headers, json=payload, timeout=30)
            response.raise_for_status()
            result = response.json()
            logging.debug(f"Stealthify response for chunk {idx+1}: {result}")
            chunk_rephrased = result.get("result")
            if chunk_rephrased:
                rephrased_text += chunk_rephrased
            else:
                logging.warning(f"No 'result' key in response for chunk {idx+1}. Using original chunk.")
                rephrased_text += chunk
        except requests.exceptions.RequestException as e:
            logging.error(f"StealthGPT API request failed for chunk {idx+1}: {str(e)}")
            if hasattr(e, 'response') and e.response is not None:
                logging.error(f"Response content: {e.response.text}")
            rephrased_text += chunk  

    logging.info("Content stealthification completed.")
    return rephrased_text

def analyze_content_with_langchain(vector_store: FAISS, original_question: str, target_word_count: int) -> dict:
    analysis_prompt = PromptTemplate(
        input_variables=["question", "content", "target_word_count"],
        template="""
        Analyze the following academic content for a paper on the question: {question}
        
        Content:
        {content}
        
        Target word count: {target_word_count}
        
        Please provide an analysis covering the following aspects:
        1. Word count: Is it close to the target? If not, by how much does it differ?
        2. Repetition: Are there any significant repetitions or redundancies?
        3. Citation: Does the content include proper citations where needed?
        4. Tone consistency: Is the tone consistently academic throughout?
        5. Conclusion quality: Does the conclusion effectively summarize the key points?
        6. Overall coherence and flow: Does the content flow logically and coherently?
        
        Provide a summary of issues found and suggestions for improvement.
        """
    )
    
    analysis_chain = RunnableSequence(
        analysis_prompt,
        model,
        StrOutputParser(),
    )
    
    # Retrieve the full content from the vector store
    full_content = " ".join([doc.page_content for doc in vector_store.similarity_search(original_question, k=len(vector_store.index_to_docstore_id))])
    
    analysis_result = analysis_chain.invoke({
        "question": original_question, 
        "content": full_content, 
        "target_word_count": target_word_count
    })
    
    # Parse the analysis result into a structured format
    # This is a simplified parsing. You might want to use a more robust method.
    issues = {
        "word_count": False,
        "repetition": False,
        "citation": False,
        "tone": False,
        "conclusion": False,
        "coherence": False
    }
    
    for issue in issues:
        if issue in analysis_result.lower():
            issues[issue] = True
    
    return {
        "analysis": analysis_result,
        "issues": issues
    }

def review_content(content: str, original_question: str, target_word_count: int, max_attempts: int = 5) -> str:
    for attempt in range(max_attempts):
        logging.info(f"Review attempt {attempt + 1}/{max_attempts}")
        try:
            humanized_content = humanize_content(content, "Academic", "High")
            logging.info("Content humanized successfully.")
            
            # Create vector store from humanized content
            vector_store = store_content_in_vector_store(humanized_content)
            
            # Analyze content using LangChain
            analysis_result = analyze_content_with_langchain(vector_store, original_question, target_word_count)
            
            if not any(analysis_result["issues"].values()):
                logging.info("Content passed all checks.")
                return humanized_content
            
            logging.info(f"Content analysis found issues: {analysis_result['issues']}")
            logging.info(f"Detailed analysis: {analysis_result['analysis']}")
            
            regenerated_content = regenerate_content(humanized_content, analysis_result, original_question, target_word_count)
            
            # Step 5: Set the regenerated content as the new content for the next iteration
            content = regenerated_content
            
        except Exception as e:
            logging.error(f"Error during content review: {str(e)}")
    
    logging.warning(f"Reached maximum attempts ({max_attempts}). Returning best version available.")
    return content

def regenerate_content(original_content: str, analysis_result: dict, question: str, target_word_count: int) -> str:
    # Split the content into chunks
    chunk_size = 750  # Adjust this based on your model's token limit
    chunks = [original_content[i:i+chunk_size] for i in range(0, len(original_content), chunk_size)]
    
    regeneration_prompt = PromptTemplate(
        input_variables=["chunk", "chunk_index", "total_chunks", "analysis", "question", "target_word_count"],
        template="""
        You are tasked with improving a part of an academic paper on the following question:
        {question}

        This is chunk {chunk_index} of {total_chunks}.

        The original content has the following issues:
        {analysis}

        Please rewrite this part of the content, addressing these issues. Ensure the following:
        1. Maintain academic tone throughout.
        2. Include proper citations where needed.
        3. Avoid repetitions and redundancies.
        4. Maintain coherence and logical flow.
        5. Ensure this part flows well with the other parts of the paper.

        If this is the last chunk, ensure a strong conclusion that summarizes key points.

        Original content for this chunk:
        {chunk}

        Provide the improved content for this chunk:
        """
    )
    
    regeneration_chain = RunnableSequence(
        regeneration_prompt,
        model,
        StrOutputParser(),
    )
    
    improved_chunks = []
    for i, chunk in enumerate(chunks):
        improved_chunk = regeneration_chain.invoke({
            "chunk": chunk,
            "chunk_index": i + 1,
            "total_chunks": len(chunks),
            "analysis": analysis_result["analysis"],
            "question": question,
            "target_word_count": target_word_count
        })
        improved_chunks.append(improved_chunk)
    
    # Combine improved chunks
    improved_content = " ".join(improved_chunks)
    
    # Final pass to ensure coherence and word count
    final_prompt = PromptTemplate(
        input_variables=["content", "analysis", "question", "target_word_count"],
        template="""
        You are tasked with finalizing an academic paper on the following question:
        {question}

        The content has been improved in chunks. Your task is to ensure overall coherence,
        smooth transitions between chunks, and that the total word count is close to {target_word_count}.

        Original analysis of issues:
        {analysis}

        Content to finalize:
        {content}

        Please provide the final, coherent version of the paper, addressing any remaining issues:
        """
    )
    
    final_chain = RunnableSequence(
        final_prompt,
        model,
        StrOutputParser(),
    )
    
    final_content = final_chain.invoke({
        "content": improved_content,
        "analysis": analysis_result["analysis"],
        "question": question,
        "target_word_count": target_word_count
    })
    
    return final_content

def save_to_docx(content: str, file_name: str) -> Document:
    doc = Document()
    doc.add_heading('Academic Paper', 0)
    paragraphs = content.split("\n\n")
    for para in paragraphs:
        doc.add_paragraph(para)
    doc.save(file_name)
    logging.info(f"Content saved to file: {file_name}")
    return doc

def academic_writing_pipeline(question: str, word_count: int, docx_file_name: str) -> Tuple[FAISS, Document, str]:
    logging.info(f"Starting academic writing pipeline. Question: '{question}', Word count: {word_count}")
    content = generate_academic_content(question, word_count, "Make sure you generate human readable content")
    logging.info(f"Initial content generated. Length: {len(content)} characters")

    reviewed_content = review_content(content, question, word_count)
    logging.info(f"Content reviewed and possibly humanized. Final length: {len(reviewed_content)} characters")

    final_vector_store = store_content_in_vector_store(reviewed_content)
    logging.info("Content stored in vector store")

    doc = save_to_docx(reviewed_content, docx_file_name)
    logging.info(f"Document saved: {docx_file_name}")

    return final_vector_store, doc, docx_file_name

if __name__ == "__main__":
    question = "In the last few decades, the biggest environmental threat worldwide is global warming, which is an increase in Earth's temperature. What are possible solutions for this issue?"
    word_count = 3000
    docx_file_name = "global_warming.docx"
    content_store, doc, file_name = academic_writing_pipeline(question, word_count, docx_file_name)
    print(f"Academic content saved in {file_name}")
    print(f"DOOLAMERCURY150")

"""
""INSTRUCTIONS FOR CLAUDE SONNET""
For this script, I want you to implement an endpoint that can be used to call it, so that an assignment can be generated. the API
endpoint should also come with its own corresponding frontend page. 
The endpoint can only be accessed by the Admin users
"""